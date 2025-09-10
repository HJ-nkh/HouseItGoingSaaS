# -*- coding: utf-8 -*-
"""
Created on Tue Nov 12 12:59:26 2024

@author: Nicolas
"""

import os
from docx import Document
from docx.shared import Pt, Cm

# Angiv stien til mappen med .docx filer
mappe_sti = "C:/Users/Nicolas/Dropbox/!P - Projekt/P0504, Hans Broges Gade 8, 2 th., 8000 Århus C , Frederik Nairn/06 Statisk rapport/ost"

# Ønskede marginer (i centimeter)
venstre_margin = 2  # 2 cm
højre_margin = 2    # 2 cm

# Skriftstørrelser (i punkter)
normal_skriftstørrelse = 9  # 12 pt for normal tekst

overskrift_skriftstørrelser = {
    'Heading 1': 12,  # 16 pt for Overskrift 1
    'Heading 2': 11,  # 14 pt for Overskrift 2
    'Heading 3': 10,
    'Heading 4': 9,
    'Heading 5': 9,
    'Heading 6': 9,
    'Heading 7': 9,
    'Heading 8': 9,
    'Heading 9': 9,
}

for filnavn in os.listdir(mappe_sti):
    if filnavn.endswith('.docx'):
        fil_sti = os.path.join(mappe_sti, filnavn)
        print(f'Behandler {filnavn}...')
        
        # Åbn dokumentet
        doc = Document(fil_sti)
        
        # Ændr sidemargener
        for sektion in doc.sections:
            sektion.left_margin = Cm(venstre_margin)
            sektion.right_margin = Cm(højre_margin)
        
        # Ændr skriftstørrelse for normal tekst
        if 'Normal' in doc.styles:
            normal_stil = doc.styles['Normal']
            normal_stil.font.size = Pt(normal_skriftstørrelse)
        
        # Ændr skriftstørrelser for overskrifter
        for stilnavn, skriftstørrelse in overskrift_skriftstørrelser.items():
            if stilnavn in doc.styles:
                stil = doc.styles[stilnavn]
                stil.font.size = Pt(skriftstørrelse)
        
        # Fjern direkte formatering fra alle afsnit og kør
        for afsnit in doc.paragraphs:
            afsnit.style = afsnit.style  # Genanvend stilen
            for run in afsnit.runs:
                run.font.size = None  # Nulstil skriftstørrelse
                run.font.name = None  # Nulstil skriftfamilie
        
        # Fjern direkte formatering fra tabelceller
        for tabel in doc.tables:
            for række in tabel.rows:
                for celle in række.cells:
                    for afsnit in celle.paragraphs:
                        afsnit.style = afsnit.style  # Genanvend stilen
                        for run in afsnit.runs:
                            run.font.size = None
                            run.font.name = None
        
        # Gem dokumentet
        doc.save(fil_sti)
        
print('Alle dokumenter er blevet behandlet.')